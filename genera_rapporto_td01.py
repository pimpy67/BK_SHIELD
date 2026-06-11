from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Stili base ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# --- Intestazione ---
header = doc.add_heading('', level=0)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run('BK SHIELD — Servizio Gestione Licenze')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Rapporto Attività — TD-01: Setup Backend')
r.font.size = Pt(13)
r.font.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Data: 10 giugno 2026   |   Sviluppatore: Andrea Pavan   |   Versione: 1.0')

doc.add_paragraph()

# --- Helper ---
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

# =====================================================================
h1('1. Obiettivo')
body(
    'Completamento del TO-DO TD-01: configurazione dell\'ambiente di sviluppo backend '
    'per il Servizio Gestione Licenze (BK SHIELD). L\'obiettivo era ottenere un server '
    'Node.js funzionante con Express, database SQLite tramite better-sqlite3, '
    'gestione sicura delle variabili d\'ambiente e middleware di error handling standardizzato.'
)

# =====================================================================
h1('2. Installazione Node.js')

h2('2.1 Versione installata')
bullet('Runtime: ', 'Node.js ')
doc.paragraphs[-1].runs[-1].text = ''
p = doc.paragraphs[-1]
p.add_run('Node.js ').bold = False
# rebuild
p = doc.add_paragraph(style='List Bullet')
p.add_run('Node.js v22.22.3').bold = True
p.add_run(' (LTS — Long Term Support)')

bullet('npm: versione inclusa nel pacchetto Node.js')
bullet('Percorso installazione: ', 'Percorso: ')
doc.paragraphs[-1].clear()
p = doc.add_paragraph(style='List Bullet')
p.add_run('Percorso installazione: ').bold = True
p.add_run('C:\\Program Files\\nodejs\\node.exe')

h2('2.2 Verifica installazione')
code_block('node --version   →  v22.22.3')

# =====================================================================
h1('3. Struttura del Progetto')

h2('3.1 Creazione cartella backend/')
body('Creata la cartella backend/ nella root del repository BK_SHIELD con la seguente struttura:')
code_block(
    'BK_SHIELD/\n'
    '└── backend/\n'
    '    ├── src/\n'
    '    │   ├── app.js\n'
    '    │   ├── server.js\n'
    '    │   ├── config/\n'
    '    │   │   ├── env.js\n'
    '    │   │   └── db.js\n'
    '    │   └── middleware/\n'
    '    │       └── errorHandler.js\n'
    '    ├── keys/\n'
    '    │   ├── private.pem    (chiave RSA privata — non tracciata in git)\n'
    '    │   └── public.pem     (chiave RSA pubblica — non tracciata in git)\n'
    '    ├── data/              (database SQLite — non tracciata in git)\n'
    '    ├── .env               (variabili ambiente — non tracciata in git)\n'
    '    ├── .env.example       (template variabili — tracciato in git)\n'
    '    ├── .gitignore\n'
    '    └── package.json'
)

h2('3.2 Inizializzazione package.json')
body('Eseguito npm init per creare il file package.json con i metadati del progetto:')
code_block(
    'nome:    bk-shield-backend\n'
    'versione: 1.0.0\n'
    'entry:   src/server.js\n'
    'script:  start | dev (nodemon) | test (jest) | test:watch'
)

# =====================================================================
h1('4. Dipendenze Installate')

h2('4.1 Dipendenze di produzione')
deps = [
    ('express ^5.2.1', 'Framework HTTP — gestione routing e middleware'),
    ('better-sqlite3 ^12.10.0', 'Driver SQLite sincrono, ottimizzato per Node.js'),
    ('dotenv ^17.4.2', 'Caricamento variabili d\'ambiente da file .env'),
    ('jsonwebtoken ^9.0.3', 'Generazione e verifica JWT (RS256)'),
    ('bcryptjs ^3.0.3', 'Hash bcrypt per API key vendor (rounds=12)'),
    ('uuid ^14.0.0', 'Generazione request_id univoci per ogni richiesta'),
    ('node-cron ^4.2.1', 'Scheduler job periodici (sistema eventi)'),
    ('nodemailer ^8.0.11', 'Invio email (OTP, notifiche)'),
    ('axios ^1.17.0', 'Client HTTP per chiamate a ERP fornitore (O1)'),
]
for pkg, desc in deps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(pkg).bold = True
    p.add_run(f' — {desc}')

h2('4.2 Dipendenze di sviluppo')
devdeps = [
    ('jest ^30.4.2', 'Framework di test unitari e di integrazione'),
    ('supertest ^7.2.2', 'Test HTTP per gli endpoint Express'),
    ('nodemon ^3.1.14', 'Riavvio automatico del server in sviluppo'),
]
for pkg, desc in devdeps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(pkg).bold = True
    p.add_run(f' — {desc}')

# =====================================================================
h1('5. Configurazione')

h2('5.1 Variabili d\'ambiente (config/env.js)')
body(
    'Implementato un modulo di configurazione centralizzato con validazione all\'avvio. '
    'Il server si blocca con errore esplicito se mancano variabili critiche.'
)
body('Variabili obbligatorie verificate all\'avvio:')
for v in ['JWT_PRIVATE_KEY_PATH', 'JWT_PUBLIC_KEY_PATH', 'LICENSE_KEY_HMAC_SECRET', 'OFFLINE_TOKEN_AES_KEY']:
    bullet(v)

body('Parametri configurabili con valori di default:')
params = [
    ('PORT', '3000'),
    ('NODE_ENV', 'development'),
    ('DATABASE_FILE', './data/licenses.sqlite'),
    ('JWT_TTL_SECONDS', '60 s'),
    ('REFRESH_TOKEN_TTL_SECONDS', '3600 s (1 ora)'),
    ('BCRYPT_ROUNDS', '12'),
    ('OTP_TTL_MINUTES', '15 min'),
    ('OTP_MAX_ATTEMPTS', '3 tentativi'),
    ('OTP_LOCKOUT_MINUTES', '30 min'),
]
for k, v in params:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(k).bold = True
    p.add_run(f' → default: {v}')

h2('5.2 Chiavi RSA per JWT RS256 (keys/)')
body(
    'Generate una coppia di chiavi RSA (2048 bit) per la firma e verifica dei JWT '
    'secondo lo standard RS256. Le chiavi sono escluse dal repository tramite .gitignore.'
)
bullet('private.pem — chiave privata, usata dal server per firmare i token')
bullet('public.pem — chiave pubblica, distribuibile ai client per la verifica')

h2('5.3 Connessione Database (config/db.js)')
body('Connessione SQLite configurata con:')
bullet('WAL mode (Write-Ahead Logging) — maggiore concorrenza in lettura')
bullet('Foreign keys ON — integrità referenziale garantita a livello DB')
bullet('Creazione automatica della cartella data/ se non esiste')

# =====================================================================
h1('6. Implementazione Server')

h2('6.1 app.js — Applicazione Express')
body('Configurato il core dell\'applicazione con:')
bullet('Parsing JSON automatico (express.json())')
bullet('Middleware request_id: ogni richiesta riceve un UUID v4 univoco per il tracciamento')
bullet('Endpoint di health check: GET /health → { status: "ok", timestamp }')
bullet('Middleware 404 (notFound) per endpoint non definiti')
bullet('Middleware error handler globale come ultimo middleware')

h2('6.2 server.js — Entry point')
body('File di avvio minimo che importa l\'app e avvia il listener sulla porta configurata.')
code_block('http://localhost:3000')

h2('6.3 middleware/errorHandler.js — Gestione errori standardizzata')
body(
    'Tutti gli errori del sistema vengono restituiti in un formato JSON uniforme, '
    'conforme alla specifica v4 dell\'analisi tecnica:'
)
code_block(
    '{\n'
    '  "error_code": "CODICE_ERRORE",\n'
    '  "message":    "Descrizione leggibile",\n'
    '  "details":    null | { ... },\n'
    '  "timestamp":  "2026-06-10T...",\n'
    '  "request_id": "uuid-v4"\n'
    '}'
)
body('Il middleware notFound genera automaticamente l\'errore ENDPOINT_NOT_FOUND (HTTP 404) per le route non definite.')

# =====================================================================
h1('7. Sicurezza e .gitignore')

h2('7.1 File esclusi dal repository')
body('Il file backend/.gitignore protegge i seguenti elementi:')
exclusions = [
    ('node_modules/', 'dipendenze npm — riproducibili con npm install'),
    ('.env / .env.local / .env.*.local', 'variabili d\'ambiente con segreti'),
    ('keys/', 'chiavi RSA private per JWT'),
    ('data/', 'database SQLite con dati di produzione'),
    ('*.sqlite', 'eventuali file SQLite aggiuntivi'),
    ('coverage/', 'report di copertura test Jest'),
    ('*.log', 'file di log applicativi e di nodemon'),
]
for item, reason in exclusions:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).bold = True
    p.add_run(f' — {reason}')

# =====================================================================
h1('8. Verifica Funzionamento')

body('Il server è stato avviato e verificato con successo:')
code_block('node src/server.js\n→ Server avviato su http://localhost:3000\n→ Ambiente: development')

body('Test health check:')
code_block(
    'GET http://localhost:3000/health\n'
    '→ 200 OK\n'
    '→ { "status": "ok", "timestamp": "2026-06-10T..." }'
)

body('Test endpoint inesistente:')
code_block(
    'GET http://localhost:3000/qualcosa\n'
    '→ 404 Not Found\n'
    '→ { "error_code": "ENDPOINT_NOT_FOUND", "message": "...", "request_id": "..." }'
)

# =====================================================================
h1('9. Riepilogo e Prossimi Passi')

h2('9.1 Stato TO-DO')
p = doc.add_paragraph(style='List Bullet')
p.add_run('TD-01 — Setup Node.js + Express + SQLite + Knex').bold = True
p.add_run(': ').bold = False
p.add_run('COMPLETATO ✓').bold = True

todos_next = [
    ('TD-02', 'Schema DB + migrazioni (tabelle principali v4 + vendor_general_setup + vendor_event_config)'),
    ('TD-03', 'Auth fornitore JWT — endpoint F1, F2'),
    ('TD-04', 'Registrazione prodotti — endpoint F6'),
]
for td, desc in todos_next:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{td} — {desc}').italic = True

h2('9.2 Note tecniche')
bullet(
    'Il progetto usa Express v5 (non Fastify come inizialmente raccomandato) — '
    'scelta equivalente per le funzionalità richieste.'
)
bullet(
    'better-sqlite3 usato direttamente senza Knex — le migrazioni saranno gestite '
    'con script SQL nativi in TD-02.'
)
bullet(
    'Il template .env.example è tracciato in git come riferimento per il setup '
    'dell\'ambiente — non contiene segreti reali.'
)

# =====================================================================
# Footer
doc.add_paragraph()
hr = doc.add_paragraph('─' * 80)
hr.runs[0].font.size = Pt(8)
hr.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    f'BK SHIELD — Documento generato il 10/06/2026 | Sviluppatore: Andrea Pavan | Revisore: Alvise'
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# =====================================================================
out_path = r'C:\Users\andrea.pavan\Documents\BK_SHIELD\Rapporto_TD01_Setup_Backend_20260610.docx'
doc.save(out_path)
print(f'Documento salvato: {out_path}')
