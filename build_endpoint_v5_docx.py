"""
Genera Endpoint_Servizio_Gestione_Licenze_v5.docx
Ereditando stili grafici da Analisi_Servizio_Gestione_Licenze_v4_correzioni.docx
"""
import re, shutil, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

BASE = r'C:\Users\andrea.pavan\Documents\BK_SHIELD'
SRC  = fr'{BASE}\Analisi_Servizio_Gestione_Licenze_v4_correzioni.docx'
DST  = fr'{BASE}\Endpoint_Servizio_Gestione_Licenze_v5.docx'
MD   = fr'{BASE}\Endpoint_Servizio_Gestione_Licenze_v5.md'

shutil.copy(SRC, DST)
doc = Document(DST)

# Svuota il body mantenendo sectPr (impostazioni pagina)
body = doc.element.body
for elem in list(body):
    if elem.tag.split('}')[-1] != 'sectPr':
        body.remove(elem)

# ── Helpers ──────────────────────────────────────────────────────────────────

def shd_elem(fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    return shd

def add_inline(para, text):
    """Converte **bold**, `code` e testo normale in runs python-docx."""
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)'
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xBE, 0x18, 0x5D)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)

def add_hr(doc):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'BFDBFE')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_code_block(doc, code_lines):
    """Blocco codice: sfondo scuro, font monospace, unico paragrafo con \\n."""
    text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd_elem('1E293B'))
    # bordo sinistro blu
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '3B82F6')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(7.5)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd_elem('EFF6FF'))
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '3B82F6')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

def add_md_table(doc, rows):
    """rows[0] = intestazione, rows[1:] = dati."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    try:
        tbl = doc.add_table(rows=len(rows), cols=n_cols)
        tbl.style = 'Table Grid'
    except Exception:
        tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.autofit = True
    for r_i, row in enumerate(rows):
        for c_i in range(n_cols):
            cell_text = row[c_i] if c_i < len(row) else ''
            cell = tbl.cell(r_i, c_i)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if r_i == 0:
                run = p.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_pr.append(shd_elem('1E40AF'))
            else:
                p.style = doc.styles['Normal']
                add_inline(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)

# ── Parser markdown ───────────────────────────────────────────────────────────

with open(MD, encoding='utf-8') as f:
    lines = [l.rstrip('\n') for l in f.readlines()]

i = 0
while i < len(lines):
    line = lines[i]

    # Riga vuota
    if not line.strip():
        i += 1
        continue

    # Heading
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = re.sub(r'\*+([^*]+)\*+', r'\1', m.group(2).strip())
        # Rimuovi markup inline dai titoli
        text  = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_heading(text, level=level)
        i += 1
        continue

    # Separatore orizzontale
    if re.match(r'^-{3,}$', line.strip()):
        add_hr(doc)
        i += 1
        continue

    # Blocco codice ```
    if line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # salta ``` di chiusura
        add_code_block(doc, code_lines)
        continue

    # Blockquote >
    if line.startswith('>'):
        bq = []
        while i < len(lines) and lines[i].startswith('>'):
            bq.append(lines[i][1:].strip())
            i += 1
        add_blockquote(doc, ' '.join(bq))
        continue

    # Tabella markdown |...|
    if line.startswith('|'):
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = lines[i].strip()
            if re.match(r'^\|[\s\-:|]+\|$', row):  # riga separatore
                i += 1
                continue
            cells = [c.strip() for c in row.split('|')[1:-1]]
            rows.append(cells)
            i += 1
        add_md_table(doc, rows)
        continue

    # Lista puntata
    m = re.match(r'^[-*]\s+(.*)', line)
    if m:
        text = m.group(1).strip()
        try:
            p = doc.add_paragraph(style='List Bullet')
        except Exception:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
        add_inline(p, text)
        i += 1
        continue

    # Lista numerata
    m = re.match(r'^\d+\.\s+(.*)', line)
    if m:
        text = m.group(1).strip()
        try:
            p = doc.add_paragraph(style='List Number')
        except Exception:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
        add_inline(p, text)
        i += 1
        continue

    # Paragrafo normale
    p = doc.add_paragraph(style='Normal')
    add_inline(p, line)
    i += 1

doc.save(DST)
sz = __import__('os').path.getsize(DST) // 1024
print(f'Generato: Endpoint_Servizio_Gestione_Licenze_v5.docx ({sz} KB)')
